"""
InsightResume AI
File Extraction Utility

Supports:
- PDF
- DOCX

Author: Aman Kumar
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import BinaryIO

import docx
import fitz  # PyMuPDF

# ==========================================================
# Logging
# ==========================================================

logger = logging.getLogger(__name__)

# ==========================================================
# File Extractor
# ==========================================================


class FileExtractor:
    """
    Handles resume validation and text extraction.

    Supported formats:
        • PDF
        • DOCX
    """

    ALLOWED_EXTENSIONS = {
        "pdf",
        "docx"
    }

    MAX_FILE_SIZE_MB = 10

    MINIMUM_TEXT_LENGTH = 100

    # ------------------------------------------------------
    # Validation
    # ------------------------------------------------------

    @classmethod
    def allowed_file(cls, filename: str) -> bool:
        """
        Check whether the uploaded file extension
        is supported.
        """

        if not filename:

            return False

        extension = filename.rsplit(".", 1)[-1].lower()

        return extension in cls.ALLOWED_EXTENSIONS

    # ------------------------------------------------------

    @classmethod
    def validate_file_size(
        cls,
        file,
        max_size_mb: int | None = None
    ) -> bool:
        """
        Validate uploaded file size.
        """

        if max_size_mb is None:

            max_size_mb = cls.MAX_FILE_SIZE_MB

        limit = max_size_mb * 1024 * 1024

        file.seek(0, 2)

        size = file.tell()

        file.seek(0)

        return size <= limit

    # ------------------------------------------------------

    @classmethod
    def get_extension(
        cls,
        filename: str
    ) -> str:

        return filename.rsplit(".", 1)[-1].lower()

    # ==========================================================
    # Public API
    # ==========================================================

    @classmethod
    def extract_text(
        cls,
        uploaded_file
    ) -> str:
        """
        Main extraction entry point.
        """

        filename = uploaded_file.filename

        extension = cls.get_extension(
            filename
        )

        logger.info(

            "Extracting text from %s",

            filename

        )

        if extension == "pdf":

            text = cls._extract_pdf(

                uploaded_file

            )

        elif extension == "docx":

            text = cls._extract_docx(

                uploaded_file

            )

        else:

            raise ValueError(

                "Unsupported file format."

            )

        text = cls.clean_text(

            text

        )

        cls.validate_text(

            text

        )

        logger.info(

            "Extraction completed (%s characters)",

            len(text)

        )

        return text

    # ==========================================================
    # Text Validation
    # ==========================================================

    @classmethod
    def validate_text(
        cls,
        text: str
    ) -> bool:

        if not text:

            raise ValueError(

                "No readable text found."

            )

        if len(text) < cls.MINIMUM_TEXT_LENGTH:

            raise ValueError(

                "Resume contains too little readable text."
            )

        return True

    # ==========================================================
    # Text Cleaning
    # ==========================================================

    @classmethod
    def clean_text(
        cls,
        text: str
    ) -> str:
        """
        Normalize extracted text.
        """

        if not text:

            return ""

        # Windows line endings

        text = text.replace(

            "\r\n",

            "\n"

        )

        # Mac line endings

        text = text.replace(

            "\r",

            "\n"

        )

        # Tabs

        text = text.replace(

            "\t",

            " "

        )

        # Multiple spaces

        text = re.sub(

            r"[ ]{2,}",

            " ",

            text

        )

        # Multiple blank lines

        text = re.sub(

            r"\n{3,}",

            "\n\n",

            text

        )

        # Remove non-printable characters

        text = "".join(

            c

            for c in text

            if c.isprintable()

            or c == "\n"

        )

        return text.strip()

    # ==========================================================
    # PDF Extraction
    # ==========================================================

    @classmethod
    def _extract_pdf(
        cls,
        uploaded_file
    ) -> str:
        """
        Extract text from a PDF using PyMuPDF.

        Supports:
        - Multi-page PDFs
        - Digital PDFs
        """

        try:

            uploaded_file.seek(0)

            pdf_bytes = uploaded_file.read()

            if not pdf_bytes:

                raise ValueError(
                    "Uploaded PDF is empty."
                )

            document = fitz.open(
                stream=pdf_bytes,
                filetype="pdf"
            )

            text_parts = []

            for page_number, page in enumerate(document, start=1):

                try:

                    page_text = page.get_text("text")

                    if page_text:

                        text_parts.append(page_text)

                except Exception:

                    logger.exception(

                        "Unable to read PDF page %s",

                        page_number

                    )

            document.close()

            extracted_text = "\n".join(text_parts)

            if not extracted_text.strip():

                raise ValueError(
                    "No readable text found in the PDF. "
                    "The file may be scanned or image-based."
                )

            logger.info(

                "PDF extraction completed (%s pages)",

                len(text_parts)

            )

            return extracted_text

        except Exception:

            logger.exception(
                "PDF extraction failed."
            )

            raise

    # ==========================================================
    # DOCX Extraction
    # ==========================================================

    @classmethod
    def _extract_docx(
        cls,
        uploaded_file
    ) -> str:
        """
        Extract text from Microsoft Word (.docx).
        """

        try:

            uploaded_file.seek(0)

            document = docx.Document(uploaded_file)

            paragraphs = []

            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:

                    paragraphs.append(text)

            # ---------- Tables ----------

            for table in document.tables:

                for row in table.rows:

                    row_values = []

                    for cell in row.cells:

                        value = cell.text.strip()

                        if value:

                            row_values.append(value)

                    if row_values:

                        paragraphs.append(
                            " | ".join(row_values)
                        )

            extracted_text = "\n".join(paragraphs)

            if not extracted_text.strip():

                raise ValueError(
                    "No readable text found in the DOCX file."
                )

            logger.info(

                "DOCX extraction completed."

            )

            return extracted_text

        except Exception:

            logger.exception(
                "DOCX extraction failed."
            )

            raise

    # ==========================================================
    # Metadata
    # ==========================================================

    @classmethod
    def get_file_metadata(
        cls,
        uploaded_file
    ) -> dict:
        """
        Return basic metadata about the uploaded file.
        """

        uploaded_file.seek(0, 2)

        size = uploaded_file.tell()

        uploaded_file.seek(0)

        filename = uploaded_file.filename

        extension = cls.get_extension(
            filename
        )

        return {

            "filename": filename,

            "extension": extension,

            "size_bytes": size,

            "size_mb": round(
                size / (1024 * 1024),
                2
            )

        }

    # ==========================================================
    # Utility Methods
    # ==========================================================

    @classmethod
    def is_pdf(
        cls,
        filename: str
    ) -> bool:

        return cls.get_extension(filename) == "pdf"

    # ----------------------------------------------------------

    @classmethod
    def is_docx(
        cls,
        filename: str
    ) -> bool:

        return cls.get_extension(filename) == "docx"

    # ----------------------------------------------------------

    @classmethod
    def get_supported_formats(cls):

        return sorted(cls.ALLOWED_EXTENSIONS)

    # ----------------------------------------------------------

    @classmethod
    def get_statistics(
        cls,
        text: str
    ) -> dict:
        """
        Return useful statistics about extracted text.
        """

        words = text.split()

        return {

            "characters": len(text),

            "words": len(words),

            "lines": len(text.splitlines()),

            "estimated_reading_time_minutes":
                max(1, round(len(words) / 200))

        }

    # ==========================================================
    # OCR Placeholder
    # ==========================================================

    @classmethod
    def extract_using_ocr(
        cls,
        uploaded_file
    ) -> str:
        """
        Reserved for future OCR support.

        This method can later integrate:
            • pytesseract
            • Google Vision
            • Azure OCR
            • AWS Textract

        Currently not implemented.
        """

        raise NotImplementedError(

            "OCR extraction is not implemented."

        )

    # ==========================================================
    # Health Check
    # ==========================================================

    @classmethod
    def health_check(cls) -> bool:
        """
        Verify required dependencies are available.
        """

        try:

            import fitz
            import docx

            return True

        except Exception:

            logger.exception(

                "FileExtractor health check failed."

            )

            return False

    # ==========================================================
    # Status
    # ==========================================================

    @classmethod
    def get_status(cls) -> dict:

        return {

            "service": "File Extractor",

            "supported_formats": cls.get_supported_formats(),

            "max_file_size_mb": cls.MAX_FILE_SIZE_MB,

            "minimum_text_length": cls.MINIMUM_TEXT_LENGTH,

            "healthy": cls.health_check()

        }

    # ==========================================================
    # Version Information
    # ==========================================================

    VERSION = "2.0.0"

    @classmethod
    def version(cls):

        return cls.VERSION

